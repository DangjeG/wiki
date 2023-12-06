import docx
import spire
from spire.doc import *


async def convert3(input_file, output_file):
    document = spire.doc.Document()

    document.LoadFromFile(input_file, FileFormat.Html, XHTMLValidationType.none)

    document.SaveToFile(output_file, FileFormat.Docx2016)
    document.Close()

    await remove_watermark(output_file, docx.Document(output_file))


async def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


async def remove_watermark(file_name: str, doc):
    i = 0
    while i < len(doc.paragraphs):
        if 'Evaluation Warning: The document was created with Spire.Doc for Python.' in doc.paragraphs[i].text:
            for j in range(1):
                await delete_paragraph(doc.paragraphs[j])
        i += 1
    doc.save(file_name)


async def export_document(file_name: str, list_document_content: list[str]):
    html = """<!DOCTYPE HTML PUBLIC "-//W3C//DTD HTML 4.01//EN" "http://www.w3.org/TR/html4/strict.dtd">
        <html>
         <head>
          <meta http-equiv="Content-Type" content="text/html; charset=utf-8">
          <title>Пример веб-страницы</title>
         </head>
         <body>"""

    for content in list_document_content:
        html += content

    html += """</body>
        </html>"""

    os.mkdir("temp_export")

    with open("temp_export/" + file_name + ".html", "w", encoding="utf-8") as file:
        file.write(html)

    await convert3("temp_export/" + file_name + ".html", "temp_export/" + file_name + ".docx")
